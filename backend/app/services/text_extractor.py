import os
from pypdf import PdfReader
import docx
from fastapi import HTTPException

class TextExtractorService:
    @staticmethod
    def extract_text(file_path: str, file_type: str) -> str:
        if not os.path.exists(file_path):
            raise HTTPException(status_code=404, detail="Uploaded file not found on server.")

        ext = file_type.lower()
        if ext == ".pdf":
            return TextExtractorService._extract_from_pdf(file_path)
        elif ext == ".docx":
            return TextExtractorService._extract_from_docx(file_path)
        else:
            raise HTTPException(status_code=400, detail=f"Unsupported extension for text extraction: {ext}")

    @staticmethod
    def _extract_from_pdf(file_path: str) -> str:
        try:
            reader = PdfReader(file_path)
            extracted_pages = []
            for i, page in enumerate(reader.pages):
                text = page.extract_text()
                if text:
                    extracted_pages.append(text.strip())
            full_text = "\n\n".join(extracted_pages)
            if not full_text.strip():
                raise HTTPException(status_code=400, detail="Could not extract text from PDF. It might be scanned/image-only.")
            return full_text
        except Exception as e:
            if isinstance(e, HTTPException):
                raise e
            raise HTTPException(status_code=500, detail=f"Failed to read PDF file: {str(e)}")

    @staticmethod
    def _extract_from_docx(file_path: str) -> str:
        try:
            doc = docx.Document(file_path)
            full_text = []
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text.strip())
            for table in doc.tables:
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_data:
                        full_text.append(" | ".join(row_data))
            result = "\n".join(full_text)
            if not result.strip():
                raise HTTPException(status_code=400, detail="DOCX file appears to be empty.")
            return result
        except Exception as e:
            if isinstance(e, HTTPException):
                raise e
            raise HTTPException(status_code=500, detail=f"Failed to read DOCX file: {str(e)}")

text_extractor_service = TextExtractorService()
