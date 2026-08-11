import os
import pytest
from app.services.text_extractor import TextExtractorService
from docx import Document
from pypdf import PdfWriter

def test_extract_text_docx(tmp_path):
    docx_path = tmp_path / "sample.docx"
    doc = Document()
    doc.add_heading("John Doe", 0)
    doc.add_paragraph("Software Engineer with 4 years experience in Python and FastAPI.")
    doc.save(str(docx_path))

    extracted = TextExtractorService.extract_text(str(docx_path), ".docx")
    assert "John Doe" in extracted
    assert "Software Engineer" in extracted
    assert "Python" in extracted

def test_extract_text_invalid_format(tmp_path):
    txt_path = tmp_path / "sample.txt"
    txt_path.write_text("Hello World")
    with pytest.raises(Exception):
        TextExtractorService.extract_text(str(txt_path), ".txt")
